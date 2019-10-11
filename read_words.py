#! python
# -*- encoding: utf-8 -*-
'''
@File    :   read_words.py
@Time    :   2019/09/24 09:30:13
@Author  :   Shen Guoxin 
@Version :   1.0
@Contact :   shenguoxin@bupt.edu.cn
@License :   (C)Copyright 2019-2020, Shenguoxin
@Desc    :   None
'''

# here put the import lib

import docx
import zipfile
import os
import shutil
'''
从word文档中读取文本和图片
'''


class ReadWords(object):
    def __init__(self):
        path = './processed_question'
        self.all_file = []
        for root, dirs, files in os.walk(path):
            self.all_file.extend(files)
        self.questions = [q[:-5] for q in self.all_file]
        self.contents = [self.get_txt_from_word(os.path.join(path, t))
                         for t in self.all_file]

    def get_txt_from_word(self, filename):
        file = docx.Document(filename)
        txts = []
        for i in range(len(file.paragraphs)):
            if len(file.paragraphs[i].text.replace(' ', '')) > 4:
                txts.append(file.paragraphs[i].text)
            else:
                txts.append('')
        return txts

    '''读取word中的table'''

    def get_table_from_word(self, filename):
        doc = docx.Document(filename)
        for table in doc.tables:  # 遍历所有表格
            print('----table------')
            for row in table.rows:  # 遍历表格的所有行
                # row_str = '\t'.join([cell.text for cell in row.cells])  # 一行数据
                # print row_str
                for cell in row.cells:
                    print(cell.text, '\t')

    def get_pic_from_word(self, path):
        store_path = 'static/tempImages/'
        zip_path = 'log.zip'
        tmp_path = 'tmp'
        os.rename(path, zip_path)
        # 进行解压
        f = zipfile.ZipFile(zip_path, 'r')
        # 将图片提取并保存
        for file in f.namelist():
            f.extract(file, tmp_path)
        # 释放该zip文件
        f.close()
        '''=============将docx文件从zip还原为docx===================='''
        os.rename(zip_path, path)
        # 得到缓存文件夹中图片列表
        pic = os.listdir(os.path.join(tmp_path, './word/media'))
        '''=============删除缓冲文件夹中的图片，用以存储下一次的文件'''
        for i in os.listdir("./tmp/word/media"):
            shutil.copy(os.path.join("./tmp/word/media", i),
                        os.path.join(store_path, i))
            os.remove(os.path.join("./tmp/word/media", i))
        return pic


if __name__ == '__main__':
    pass
